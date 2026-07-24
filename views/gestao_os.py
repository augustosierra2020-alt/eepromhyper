import streamlit as st
import pandas as pd
import io
import os
import re
import time
import sqlite3
from datetime import datetime
from services.hf_sync import backup_local_para_nuvem_async
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BASE_DIR = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))

# ==========================================
# CONEXÃO ISOLADA E SEGURA
# ==========================================
def conectar_db_seguro():
    """Garante uma conexão nova e exclusiva para evitar o erro de 'closed database'"""
    db_path = os.path.join(BASE_DIR, "eeprom_master.db")
    conn = sqlite3.connect(db_path, timeout=30.0)
    conn.execute("PRAGMA foreign_keys = ON;")
    return conn

# ==========================================
# REGRAS DE ENGENHARIA AUTOMOTIVA
# ==========================================
def calcular_valor_inicial(linha):
    descricao = str(linha.get("Nome arquivo", linha.get("Descrição", ""))).upper().strip()
    veiculo = str(linha.get("Fabricante", linha.get("Veículo", ""))).upper().strip()
    
    if "TOTAL" in descricao or "TOTAL" in veiculo: return 0.0
        
    eh_especial = any(fab in veiculo for fab in ["NEW HOLLAND", "VALTRA", "CASE", "MASSEY", "CLAAS", "DEERE", "FENDT", "JACTO", "VOLVO CE"]) and "VOLVO TRUCK" not in veiculo
    has_p420 = bool(re.search(r'P\s*0?\s*420', descricao))
    
    if "MOD" in descricao and has_p420: return 1400 if eh_especial else 650
    if has_p420: return 200
    if re.search(r'S+T+A*G+E*\s*2', descricao): return 1400 if eh_especial else 650
    elif "MOD" in descricao or "OFF" in descricao: return 700 if eh_especial else 350
    
    return linha.get("Valor", None) if pd.notna(linha.get("Valor", None)) and str(linha.get("Valor", None)).strip() != "" else None

def limpar_descricao_os(desc):
    desc = str(desc).upper().strip()
    match = re.search(r'P\s*0?\s*420', desc)
    if match: 
        return f"MOD OFF {match.group(0).replace(' ', '')}" if "MOD" in desc and "OFF" in desc else f"MOD {match.group(0).replace(' ', '')}" if "MOD" in desc else f"OFF {match.group(0).replace(' ', '')}" if "OFF" in desc else match.group(0).replace(" ", "")
    if re.search(r'S+T+A*G+E*\s*2', desc): return "STAG 2"
    return "MOD" if "MOD" in desc else "OFF" if "OFF" in desc else desc

def higienizar_valor_monetario_para_calculo(val):
    if pd.isna(val) or str(val).strip() == "": return 0.0
    val = str(val).upper().replace("R$", "").strip()
    try: 
        return float(val.replace(".", "").replace(",", ".") if "." in val and "," in val else val.replace(",", "."))
    except: 
        return 0.0

def modificar_modelo_docx(modelo_bytes, flash_point, cliente_nome, city, contato, linhas_tabela, total_valor):
    doc = Document(io.BytesIO(modelo_bytes))
    for t in doc.tables:
        for row in t.rows:
            if len(row.cells) >= 2:
                txt = row.cells[0].text.upper().strip()
                if "CLIENTE:" in txt: row.cells[1].text = f"{cliente_nome} - {flash_point}"
                elif "CIDADE:" in txt: row.cells[1].text = city
                elif "CONTATO:" in txt: row.cells[1].text = contato

    linhas_validas = [l for l in linhas_tabela if pd.notna(l.get("Valor")) and str(l.get("Valor")).strip() not in ["", "NAN", "0", "0.0"]]
    tabela_servicos = next((t for t in doc.tables if len(t.rows) > 0 and "Nº MAPA" in t.rows[0].cells[0].text.upper()), None)
            
    if tabela_servicos:
        for i, linha in enumerate(linhas_validas):
            idx = i + 1  
            row_cells = tabela_servicos.add_row().cells if idx >= len(tabela_servicos.rows) else tabela_servicos.rows[idx].cells
            val_fmt = f"R$ {higienizar_valor_monetario_para_calculo(linha.get('Valor', 0)):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            dados = [str(linha.get("Nº Mapa", "")), str(linha.get("Data", "")), str(linha.get("Veículo", "")), str(linha.get("Placa", "")), limpar_descricao_os(linha.get("Descrição", "")), val_fmt]
            
            for c_idx, val in enumerate(dados):
                if c_idx < len(row_cells):
                    row_cells[c_idx].text = val
                    for p in row_cells[c_idx].paragraphs:
                        if c_idx in [0, 1, 3, 5]: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs: r.font.name = 'Arial'; r.font.size = Pt(10)
        while len(tabela_servicos.rows) > len(linhas_validas) + 1: 
            tabela_servicos._tbl.remove(tabela_servicos.rows[len(linhas_validas) + 1]._tr)

    val_tot_fmt = f"{float(total_valor if pd.notna(total_valor) else 0.0):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    for t in doc.tables:
        for row in t.rows:
            if any("TOTAL" in c.text.upper() for c in row.cells):
                for c in row.cells:
                    if "R$" in c.text or "NAN" in c.text.upper() or c.text.strip() == "":
                        c.text = f"R$ {val_tot_fmt}"
                        for p in c.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for r in p.runs: r.bold = True; r.font.name = 'Arial'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(234, 88, 12)
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

# ==========================================
# RENDERIZAÇÃO INTERFACE GESTÃO E OS
# ==========================================
def render_gestao_os():
    st.title("📊 Gestão Estrutural & Ordens de Serviço (OS)")
    
    # Inicializadores de Estado Protegidos
    if "df_filtrado" not in st.session_state: st.session_state.df_filtrado = None
    if "os_mes_selecionado" not in st.session_state: st.session_state.os_mes_selecionado = ""
    if "os_cliente_selecionado" not in st.session_state: st.session_state.os_cliente_selecionado = ""
    if "backup_cliente_sel" not in st.session_state: st.session_state.backup_cliente_sel = ""

    aba1, aba2, aba3_mensal, aba3_geral, aba4 = st.tabs([
        "📋 Processamento", 
        "📄 Gerar OS", 
        "📅 Histórico Mensal", 
        "💾 Backup de Clientes", 
        "📊 Painel"
    ])
    
    # Abrimos a conexão de forma segura usando bloco Try/Finally
    conn = conectar_db_seguro()
    try:
        cursor = conn.cursor()

        # ====================================================
        # ABA 1: PROCESSAMENTO E EXTRAÇÃO AUTOMATIZADA
        # ====================================================
        with aba1:
            st.subheader("Processamento Inteligente (Suporte a Múltiplos Arquivos)")
            arquivos_carregados = st.file_uploader("Arraste uma ou mais planilhas de faturamento:", type=["xlsx", "xls", "csv"], accept_multiple_files=True)
            
            if arquivos_carregados:
                try:
                    dfs = []
                    for arquivo in arquivos_carregados:
                        conteudo = arquivo.read()
                        if arquivo.name.endswith('.csv'):
                            df_temp = pd.read_csv(io.BytesIO(conteudo), sep=";", encoding="utf-8")
                        else:
                            df_temp = pd.read_excel(io.BytesIO(conteudo))
                        
                        df_temp.columns = df_temp.columns.str.strip()
                        for col in list(df_temp.columns):
                            if str(col).upper().replace(" ", "") == "FLASHPOINT": df_temp.rename(columns={col: "Flash Point"}, inplace=True)
                        df_temp.rename(columns={"Arquivo ID": "Nº Mapa", "Fabricante": "Veículo", "Matrícula": "Placa", "Nome arquivo": "Descrição", "Dada": "Data"}, inplace=True, errors="ignore")
                        dfs.append(df_temp)
                    
                    if dfs:
                        df = pd.concat(dfs, ignore_index=True)
                        df = df[~df.apply(lambda row: row.astype(str).str.contains('TOTAL', case=False).any(), axis=1)]
                        
                        df_filtrado = df.copy()
                        df_filtrado["Valor"] = df_filtrado.apply(calcular_valor_inicial, axis=1)
                        df_filtrado = df_filtrado[df_filtrado["Valor"] > 0]
                        
                        st.session_state.df_filtrado = df_filtrado
                        st.dataframe(df_filtrado, use_container_width=True)
                        
                        with st.container(border=True):
                            st.markdown("### 📅 Arquivamento Definitivo no Banco de Dados Mestre")
                            col_sm1, col_sm2 = st.columns(2)
                            mes_nome_in = col_sm1.selectbox("Informe o Mês:", ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"], index=datetime.now().month-1)
                            ano_nome_in = col_sm2.number_input("Informe o Ano:", value=datetime.now().year, step=1)
                            pasta_mes_designada = f"{mes_nome_in} - {ano_nome_in}"
                            
                            if st.button("💾 Gravar e Blindar Planilha no Cofre", use_container_width=True, type="primary"):
                                tot_faturado = pd.to_numeric(df_filtrado["Valor"], errors='coerce').sum()
                                cursor.execute("INSERT OR REPLACE INTO planilhas_mensais (mes_ano, dados_json, total_faturado) VALUES (?, ?, ?)", (pasta_mes_designada, df_filtrado.to_json(orient="records"), float(tot_faturado)))
                                
                                if "Flash Point" in df_filtrado.columns:
                                    for fp_item in df_filtrado["Flash Point"].dropna().unique():
                                        fp_clean = str(fp_item).strip().upper()
                                        if fp_clean: cursor.execute("INSERT OR IGNORE INTO clientes_fp (fp_codigo, cidade, contato) VALUES (?, '', '')", (fp_clean,))
                                conn.commit()
                                backup_local_para_nuvem_async()
                                st.success(f"✅ Sucesso Absoluto! Lote de faturamento '{pasta_mes_designada}' blindado permanentemente!")
                except Exception as e: 
                    st.error(f"Erro no processamento da planilha: {e}")

        # ====================================================
        # ABA 2: EMISSOR DE OS INDIVIDUAL/LOTE VIA DOCX
        # ====================================================
        with aba2:
            st.subheader("📄 Emissor de Ordem de Serviço")
            modelo_word = st.file_uploader("Selecione o arquivo de Template .docx:", type=["docx"])
            if st.session_state.df_filtrado is None or st.session_state.df_filtrado.empty: 
                st.info("Carregue e filtre uma planilha na aba anterior primeiro.")
            elif modelo_word is None: 
                st.warning("Anexe o template Word da oficina.")
            else:
                df_base = st.session_state.df_filtrado
                fp_list = sorted(list(set(str(v).strip() for v in df_base["Flash Point"].unique() if pd.notna(v))))
                fp_sel = st.selectbox("Selecione o Código do Cliente (Flash Point):", fp_list)
                
                bloco = df_base[df_base["Flash Point"] == fp_sel]
                cliente_in = st.text_input("Nome do Cliente Responsável:", value=str(bloco.iloc[0].get("Cliente", f"Parceiro {fp_sel}")))
                
                df_prev = bloco.copy()
                if "Descrição" in df_prev.columns: df_prev["Descrição"] = df_prev["Descrição"].apply(limpar_descricao_os)
                df_ed = st.data_editor(df_prev, num_rows="dynamic", use_container_width=True)
                soma_total = df_ed["Valor"].apply(higienizar_valor_monetario_para_calculo).sum()
                
                st.metric("Valor Consolidado da OS (Serviços Puros)", f"R$ {soma_total:,.2f}")
                
                if st.button("🚀 Preencher e Gerar Documento de OS", use_container_width=True):
                    buffer_docx = modificar_modelo_docx(modelo_word.read(), fp_sel, cliente_in, "Sede Técnica", "Contato Central", df_ed.to_dict(orient="records"), soma_total)
                    st.session_state.docx_atual_buffer = buffer_docx.getvalue()
                    st.session_state.docx_atual_nome = f"{fp_sel} - OS.docx"
                    st.session_state.docx_atual_valor = soma_total
                    st.session_state.docx_atual_fp = fp_sel
                    st.success("Ordem de serviço gerada com sucesso e pronta para arquivamento!")

                if "docx_atual_buffer" in st.session_state and st.session_state.get("docx_atual_fp") == fp_sel:
                    st.download_button("📥 Baixar Arquivo da OS (.docx)", data=st.session_state.docx_atual_buffer, file_name=st.session_state.docx_atual_nome, use_container_width=True)
                    
                    st.markdown("### 💾 Opção: Salvar e Destinar à Pasta Compartilhada do Cliente")
                    col_c1, col_c2 = st.columns(2)
                    m_os = col_c1.selectbox("Mês de Competência:", ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"], index=datetime.now().month-1, key="os_mes_comp")
                    a_os = col_c2.number_input("Ano Alvo:", value=datetime.now().year, step=1, key="os_ano_comp")
                    periodo_chave_os = f"{m_os} - {a_os}"
                    
                    if st.button("💾 Armazenar Permanentemente na Pasta Virtual do Cliente", type="primary", use_container_width=True):
                        cursor.execute("INSERT OR IGNORE INTO clientes_fp (fp_codigo, cidade, contato) VALUES (?, '', '')", (str(fp_sel).strip().upper(),))
                        nome_custom = f"{fp_sel} - OS {m_os} {a_os}.docx"
                        cursor.execute("INSERT INTO os_salvas (fp_codigo, mes_ano, nome_arquivo, dados_bytes, valor_total) VALUES (?, ?, ?, ?, ?)", (str(fp_sel).strip().upper(), periodo_chave_os, nome_custom, st.session_state.docx_atual_buffer, float(st.session_state.docx_atual_valor)))
                        conn.commit()
                        backup_local_para_nuvem_async()
                        st.success(f"✅ OS blindada na biblioteca de arquivos digitais do parceiro {fp_sel}!")

        # ====================================================
        # ABA 3: PASTAS DIGITAIS DOS CLIENTES (VISÃO MENSAL)
        # ====================================================
        with aba3_mensal:
            cursor.execute("SELECT DISTINCT mes_ano FROM planilhas_mensais")
            p1 = [r[0] for r in cursor.fetchall()]
            cursor.execute("SELECT DISTINCT mes_ano FROM os_salvas")
            p2 = [r[0] for r in cursor.fetchall()]
            lista_meses_arquivados = sorted(list(set(p1 + p2)))

            if st.session_state.os_mes_selecionado == "":
                st.subheader("📅 Central Histórica: Pastas por Mês")
                st.markdown("### Selecione o Período Mensal para gerenciar as pastas dos parceiros:")
                if not lista_meses_arquivados: st.info("Nenhuma pasta ou faturamento arquivado até o momento.")
                else:
                    for i in range(0, len(lista_meses_arquivados), 4):
                        cols_m = st.columns(4)
                        for j in range(4):
                            if i + j < len(lista_meses_arquivados):
                                m_folder = lista_meses_arquivados[i + j]
                                with cols_m[j]:
                                    with st.container(border=True):
                                        st.markdown(f"<div style='display: flex; justify-content: center; align-items: center; height: 110px; width: 100%; margin-bottom: 10px; background: linear-gradient(135deg, #1E88E5 0%, #0D47A1 100%); border-radius: 12px;'><p style='text-align:center; font-weight:bold; color:#FFFFFF; margin:0; font-size: 1.1rem;'>📅 {m_folder}</p></div>", unsafe_allow_html=True)
                                        if st.button("Abrir Pasta Mensal", key=f"f_m_{m_folder}", use_container_width=True):
                                            st.session_state.os_mes_selecionado = m_folder
                                            st.session_state.os_cliente_selecionado = ""
                                            st.rerun()
            else:
                col_b1, col_t1 = st.columns([1, 4])
                if col_b1.button("⬅️ Mudar de Mês", use_container_width=True, type="primary"):
                    st.session_state.os_mes_selecionado = ""; st.session_state.os_cliente_selecionado = ""; st.rerun()
                col_t1.markdown(f"## 📂 Pasta Mensal: {st.session_state.os_mes_selecionado}")
                st.markdown("---")

                cursor.execute("SELECT DISTINCT fp_codigo FROM os_salvas WHERE mes_ano = ?", (st.session_state.os_mes_selecionado,))
                c_os = [r[0] for r in cursor.fetchall()]
                cursor.execute("SELECT dados_json FROM planilhas_mensais WHERE mes_ano = ?", (st.session_state.os_mes_selecionado,))
                row_json = cursor.fetchone()
                
                c_pl = []
                if row_json:
                    try: 
                        df_m = pd.read_json(io.StringIO(row_json[0]))
                        if "Flash Point" in df_m.columns: c_pl = df_m["Flash Point"].dropna().unique().tolist()
                    except: pass
                
                clientes_do_mes = sorted(list(set([str(c).strip().upper() for c in (c_os + c_pl) if c])))

                if st.session_state.os_cliente_selecionado == "":
                    st.markdown("### Clientes identificados neste período:")
                    if not clientes_do_mes: st.info("Nenhum parceiro registrou ordens de serviço neste período.")
                    else:
                        for i in range(0, len(clientes_do_mes), 4):
                            cols_c = st.columns(4)
                            for j in range(4):
                                if i + j < len(clientes_do_mes):
                                    c_code = clientes_do_mes[i + j]
                                    with cols_c[j]:
                                        with st.container(border=True):
                                            st.markdown(f"<div style='display: flex; justify-content: center; align-items: center; height: 90px; width: 100%; margin-bottom: 10px; background: linear-gradient(135deg, #8E24AA 0%, #4A148C 100%); border-radius: 12px;'><p style='text-align:center; font-weight:bold; color:#FFFFFF; margin:0; font-size: 1rem;'>📁 {c_code}</p></div>", unsafe_allow_html=True)
                                            if st.button(f"Ver Pasta de {c_code}", key=f"f_c_{c_code}", use_container_width=True):
                                                st.session_state.os_cliente_selecionado = c_code; st.rerun()
                else:
                    col_b2, col_t2 = st.columns([1, 4])
                    if col_b2.button("⬅️ Mudar Cliente", use_container_width=True): 
                        st.session_state.os_cliente_selecionado = ""; st.rerun()
                    col_t2.markdown(f"### 🗂️ OS Salvas no Mês para: {st.session_state.os_cliente_selecionado}")
                    
                    cursor.execute("SELECT id, mes_ano, nome_arquivo, dados_bytes, valor_total FROM os_salvas WHERE fp_codigo = ? AND mes_ano = ? ORDER BY id DESC", (st.session_state.os_cliente_selecionado, st.session_state.os_mes_selecionado))
                    arquivos_cliente = cursor.fetchall()
                    
                    if not arquivos_cliente: st.info("Nenhum documento anexado ao histórico deste cliente neste mês.")
                    else:
                        for doc_id, doc_mes, doc_nome, doc_bytes, doc_val in arquivos_cliente:
                            with st.container(border=True):
                                col_f1, col_f2, col_f3 = st.columns([4, 1, 1])
                                col_f1.write(f"📄 **{doc_nome}** \n\n📅 Referência: `{doc_mes}` | 💰 Faturamento da OS: **R$ {doc_val:,.2f}**")
                                col_f2.download_button("📥 Baixar Documento", data=doc_bytes, file_name=doc_nome, mime="application/octet-stream", key=f"dl_f_{doc_id}", use_container_width=True)
                                if col_f3.button("🗑️ Excluir", key=f"del_f_{doc_id}", use_container_width=True):
                                    cursor.execute("DELETE FROM os_salvas WHERE id = ?", (doc_id,))
                                    conn.commit()
                                    backup_local_para_nuvem_async()
                                    st.rerun()

                    st.markdown("---")
                    st.markdown("### Anexar Manualmente Documento Adicional ao Parceiro (Neste Mês)")
                    up_manual = st.file_uploader("Arraste ordens assinadas ou laudos mecânicos:", key="up_manual_shared")
                    val_m = st.number_input("Valor do Serviço Anexo (R$):", value=0.0, step=50.0)
                    
                    if st.button("💾 Inserir Documento", use_container_width=True, type="primary"):
                        if up_manual:
                            cursor.execute("INSERT INTO os_salvas (fp_codigo, mes_ano, nome_arquivo, dados_bytes, valor_total) VALUES (?, ?, ?, ?, ?)", (st.session_state.os_cliente_selecionado, st.session_state.os_mes_selecionado, up_manual.name, up_manual.read(), float(val_m)))
                            conn.commit()
                            backup_local_para_nuvem_async()
                            st.success("Documento injetado e sincronizado com a nuvem!"); st.rerun()

        # ====================================================
        # ABA 4: BACKUP DE CLIENTES (PASTA VIVA E UNIFICADA DO BANCO)
        # ====================================================
        with aba3_geral:
            st.subheader("💾 Backup de Clientes")
            st.caption("Acesse diretamente a pasta viva e centralizada de cada cliente, contendo todo o histórico de serviços independente do mês.")
            
            cursor.execute("SELECT DISTINCT fp_codigo FROM os_salvas")
            c_os_all = [r[0] for r in cursor.fetchall()]
            cursor.execute("SELECT DISTINCT fp_codigo FROM clientes_fp")
            c_fp_all = [r[0] for r in cursor.fetchall()]
            
            todos_clientes_com_os = sorted(list(set([str(c).strip().upper() for c in (c_os_all + c_fp_all) if c])))
            
            if st.session_state.backup_cliente_sel == "":
                st.markdown("### Selecione o cliente para acessar a pasta viva:")
                if not todos_clientes_com_os:
                    st.info("Nenhum cliente possui arquivos de OS salvos no banco de dados ainda.")
                else:
                    for i in range(0, len(todos_clientes_com_os), 4):
                        cols_bc = st.columns(4)
                        for j in range(4):
                            if i + j < len(todos_clientes_com_os):
                                cli = todos_clientes_com_os[i + j]
                                with cols_bc[j]:
                                    with st.container(border=True):
                                        st.markdown(f"<div style='display: flex; justify-content: center; align-items: center; height: 90px; width: 100%; margin-bottom: 10px; background: linear-gradient(135deg, #FF8C00 0%, #E65100 100%); border-radius: 12px;'><p style='text-align:center; font-weight:bold; color:#FFFFFF; margin:0; font-size: 1rem;'>👤 {cli}</p></div>", unsafe_allow_html=True)
                                        if st.button(f"Abrir Pasta de {cli}", key=f"bkp_btn_{cli}", use_container_width=True):
                                            st.session_state.backup_cliente_sel = cli
                                            st.rerun()
            else:
                cli_atual = st.session_state.backup_cliente_sel
                c1_bkp, c2_bkp = st.columns([1, 4])
                if c1_bkp.button("⬅️ Voltar à Lista", type="primary"):
                    st.session_state.backup_cliente_sel = ""
                    st.rerun()
                c2_bkp.markdown(f"## 📂 Pasta Unificada Viva: {cli_atual}")
                st.markdown("---")
                
                cursor.execute("SELECT id, mes_ano, nome_arquivo, dados_bytes, valor_total FROM os_salvas WHERE fp_codigo = ? ORDER BY id DESC", (cli_atual,))
                todo_historico_cliente = cursor.fetchall()
                
                st.markdown(f"#### Histórico Completo de Serviços Anexados ({len(todo_historico_cliente)} arquivo(s))")
                if not todo_historico_cliente:
                    st.info("Nenhuma OS salva nesta pasta viva ainda.")
                else:
                    for doc_id, doc_mes, doc_nome, doc_bytes, doc_val in todo_historico_cliente:
                        with st.container(border=True):
                            col_h1, col_h2, col_h3 = st.columns([4, 1, 1])
                            col_h1.write(f"📄 **{doc_nome}** \n\n📅 Mês Referência: `{doc_mes}` | 💰 Faturamento: **R$ {doc_val:,.2f}**")
                            col_h2.download_button("📥 Baixar", data=doc_bytes, file_name=doc_nome, mime="application/octet-stream", key=f"dl_bkp_{doc_id}", use_container_width=True)
                            if col_h3.button("🗑️ Excluir", key=f"del_bkp_{doc_id}", use_container_width=True):
                                cursor.execute("DELETE FROM os_salvas WHERE id = ?", (doc_id,))
                                conn.commit(); backup_local_para_nuvem_async(); st.rerun()

        # ====================================================
        # ABA 5: MONITORAMENTO FINANCEIRO E AUDITORIA MENSAL
        # ====================================================
        with aba4:
            st.subheader("📊 Painel de Monitoramento Mensal (Réplica de Planilhas)")
            cursor.execute("SELECT DISTINCT mes_ano FROM planilhas_mensais")
            p_sheets = [r[0] for r in cursor.fetchall()]
            cursor.execute("SELECT DISTINCT mes_ano FROM os_salvas")
            p_docs = [r[0] for r in cursor.fetchall()]
            todos_periodos = sorted(list(set(p_sheets + p_docs)))
            
            if not todos_periodos: st.info("Sem fechamentos ativos no banco remoto.")
            else:
                mes_escolhido = st.selectbox("Selecione o Período para Auditoria Financeira:", todos_periodos)
                if mes_escolhido:
                    cursor.execute("SELECT dados_json, total_faturado FROM planilhas_mensais WHERE mes_ano = ?", (mes_escolhido,))
                    row_p = cursor.fetchone()
                    cursor.execute("SELECT fp_codigo, valor_total FROM os_salvas WHERE mes_ano = ?", (mes_escolhido,))
                    rows_os = cursor.fetchall()
                    
                    tot_p = row_p[1] if row_p else 0.0
                    tot_os = sum([r[1] for r in rows_os])
                    
                    st.markdown(f"### Resumo Financeiro Consolidado: {mes_escolhido}")
                    nc1, nc2 = st.columns(2)
                    with nc1.container(border=True):
                        st.markdown("📈 **Faturamento Planilha Base (Serviços Totais)**")
                        st.markdown(f"<h2 style='color:#1E88E5;'>R$ {tot_p:,.2f}</h2>", unsafe_allow_html=True)
                    with nc2.container(border=True):
                        st.markdown("📄 **Valor Capturado em OS Individuais Emitidas**")
                        st.markdown(f"<h2 style='color:#EA580C;'>R$ {tot_os:,.2f}</h2>", unsafe_allow_html=True)
                    
                    if row_p and row_p[0]:
                        st.markdown(f"### 📋 Planilha FP Original Arquivada para Consulta Coletiva ({mes_escolhido})")
                        try:
                            df_salvo = pd.read_json(io.StringIO(row_p[0]))
                            st.dataframe(df_salvo, use_container_width=True)
                        except Exception as err:
                            st.error(f"Erro ao reestruturar planilha extraída: {err}")

                    if st.button("🗑️ Destruir Registro Mensal Deste Período (Ação Irreversível)", use_container_width=True):
                        cursor.execute("DELETE FROM planilhas_mensais WHERE mes_ano = ?", (mes_escolhido,))
                        conn.commit()
                        backup_local_para_nuvem_async()
                        st.warning("Faturamento apagado definitivamente da nuvem."); st.rerun()

    # O bloco "finally" SEMPRE executará quando a função terminar ou o app recarregar
    finally:
        if 'conn' in locals():
            conn.close()