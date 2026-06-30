import streamlit as st
import os
import json
import base64
import sqlite3
import re
import shutil
import unicodedata
import pandas as pd
import io
from PIL import Image
from openpyxl.styles import PatternFill
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from duckduckgo_search import DDGS
from huggingface_hub import HfApi, hf_hub_download, InferenceClient

# ==========================================
# 1. CONFIGURAÇÕES DE NUVEM E ANCORAGEM
# ==========================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, "eeprom_master.db")

HF_TOKEN = os.environ.get("HF_TOKEN")
DATASET_REPO_ID = "GrizzlyBear25/HyperTork_DB" 

def sincronizar_nuvem_para_local():
    if HF_TOKEN:
        try:
            db_nuvem = hf_hub_download(repo_id=DATASET_REPO_ID, filename="eeprom_master.db", repo_type="dataset", token=HF_TOKEN)
            shutil.copy(db_nuvem, DB_PATH)
            return True
        except Exception:
            pass
    return False

def backup_local_para_nuvem():
    if HF_TOKEN and os.path.exists(DB_PATH):
        try:
            api = HfApi()
            api.upload_file(
                path_or_fileobj=DB_PATH,
                path_in_repo="eeprom_master.db",
                repo_id=DATASET_REPO_ID,
                repo_type="dataset",
                token=HF_TOKEN
            )
            return True
        except Exception:
            pass
    return False

def higienizar_nome(nome):
    if not nome: return ""
    nome_limpo = " ".join(nome.strip().upper().split())
    return re.sub(r'[\\/*?:"<>|]', "", nome_limpo)

def limpar_para_comparacao(texto):
    if not texto: return ""
    texto = unicodedata.normalize('NFKD', str(texto)).encode('ASCII', 'ignore').decode('utf-8')
    return re.sub(r'[^A-Z0-9]', '', texto.upper())

def sincronizar_banco_com_pastas():
    conn = conectar_db()
    cursor = conn.cursor()
    cursor.execute("SELECT nome FROM montadoras")
    montadoras = cursor.fetchall()
    for m in montadoras:
        os.makedirs(os.path.join(BASE_DIR, m[0]), exist_ok=True)
    
    cursor.execute("SELECT montadora_nome, modelo, posicao_inicio, intervalo, valores_invertidos, escala, detalhes, id FROM veiculos")
    veiculos = cursor.fetchall()
    
    pastas_criadas = 0
    for v in veiculos:
        mont, mod, ini, inter, val_inv, esc, det, v_id = v
        pasta_alvo = os.path.join(BASE_DIR, mont, mod)
        if not os.path.exists(pasta_alvo):
            os.makedirs(pasta_alvo)
            pastas_criadas += 1
            dados_json = {"posicao_inicio": ini, "intervalo": inter, "valores_invertidos": val_inv, "escala": esc, "detalhes": det}
            with open(os.path.join(pasta_alvo, "dados.json"), "w", encoding="utf-8") as f:
                json.dump(dados_json, f, indent=4, ensure_ascii=False)
            cursor.execute("SELECT foto, ordem FROM graficos WHERE veiculo_id = ?", (v_id,))
            for foto_blob, ordem in cursor.fetchall():
                with open(os.path.join(pasta_alvo, f"grafico_{ordem}.png"), "wb") as img_f:
                    img_f.write(foto_blob)
    conn.close()
    return pastas_criadas

# ==========================================
# 2. SETUP DO BANCO DE DADOS UNIFICADO
# ==========================================
def conectar_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("PRAGMA foreign_keys = ON;")
    return conn

def init_db():
    conn = conectar_db()
    cursor = conn.cursor()
    cursor.execute("CREATE TABLE IF NOT EXISTS montadoras (id INTEGER PRIMARY KEY AUTOINCREMENT, nome TEXT UNIQUE NOT NULL)")
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS veiculos (
            id INTEGER PRIMARY KEY AUTOINCREMENT, montadora_nome TEXT NOT NULL, modelo TEXT NOT NULL,
            posicao_inicio TEXT, intervalo TEXT, valores_invertidos TEXT, escala TEXT, detalhes TEXT,
            UNIQUE(montadora_nome, modelo)
        )
    """)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS graficos (
            id INTEGER PRIMARY KEY AUTOINCREMENT, veiculo_id INTEGER NOT NULL, foto BLOB NOT NULL, ordem INTEGER NOT NULL,
            FOREIGN KEY (veiculo_id) REFERENCES veiculos(id) ON DELETE CASCADE
        )
    """)
    cursor.execute("CREATE TABLE IF NOT EXISTS chip_memoria (id INTEGER PRIMARY KEY AUTOINCREMENT, chave TEXT UNIQUE NOT NULL, valor TEXT NOT NULL)")
    
    cursor.execute('''CREATE TABLE IF NOT EXISTS obd2_history
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  codigo TEXT, montadora TEXT, modelo TEXT, ano TEXT,
                  descricao TEXT, data TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
                  
    try: cursor.execute("ALTER TABLE obd2_history ADD COLUMN montadora TEXT")
    except: pass
    try: cursor.execute("ALTER TABLE obd2_history ADD COLUMN modelo TEXT")
    except: pass
    try: cursor.execute("ALTER TABLE obd2_history ADD COLUMN ano TEXT")
    except: pass
    try: cursor.execute("ALTER TABLE obd2_history ADD COLUMN segmento TEXT")
    except: pass
    
    conn.commit()
    conn.close()

sincronizar_nuvem_para_local()
init_db()
sincronizar_banco_com_pastas()

def mapear_pasta_logos(base):
    for d in os.listdir(base):
        if d.lower() in ['logos', 'logo'] and os.path.isdir(os.path.join(base, d)):
            return os.path.join(base, d)
    return os.path.join(base, "Logos")

LOGOS_DIR = mapear_pasta_logos(BASE_DIR)
if not os.path.exists(LOGOS_DIR):
    os.makedirs(LOGOS_DIR)

st.set_page_config(page_title="HyperTork System Hub", page_icon="⚙️", layout="wide")

# ==========================================
# 3. ROTINAS AUXILIARES DE IMAGEM (LOGOS)
# ==========================================
@st.cache_data(show_spinner=False)
def buscar_logo_montadora_automatica(montadora):
    if os.path.exists(LOGOS_DIR):
        arquivos = sorted(os.listdir(LOGOS_DIR), key=lambda x: (not x.lower().endswith('.png'), x))
        mont_alvo = limpar_para_comparacao(montadora)
        for arquivo in arquivos:
            if not arquivo.lower().endswith(('.png', '.jpg', '.jpeg', '.webp', '.gif')): continue
            nome_base = os.path.splitext(arquivo)[0]
            nome_arq = limpar_para_comparacao(nome_base)
            if mont_alvo == nome_arq: return os.path.join(LOGOS_DIR, arquivo)
        for arquivo in arquivos:
            if not arquivo.lower().endswith(('.png', '.jpg', '.jpeg', '.webp', '.gif')): continue
            nome_base = os.path.splitext(arquivo)[0]
            nome_arq = limpar_para_comparacao(nome_base)
            if mont_alvo in nome_arq or nome_arq in mont_alvo: return os.path.join(LOGOS_DIR, arquivo)
    return None

@st.cache_data(show_spinner=False)
def obter_image_base64_html(caminho):
    try:
        extensao = caminho.split('.')[-1].lower()
        mime = "image/png" if extensao == 'png' else ("image/jpeg" if extensao in ['jpg', 'jpeg'] else f"image/{extensao}")
        with open(caminho, "rb") as image_file:
            encoded = base64.b64encode(image_file.read()).decode()
            return f"data:{mime};base64,{encoded}"
    except: return ""

# ==========================================
# 4. ROTINAS OBD-II / DTC DE ALTA PERFORMANCE (RAG UNIVERSAL)
# ==========================================
def salvar_pesquisa_obd2(codigo, segmento, montadora, modelo, ano, descricao):
    conn = conectar_db()
    c = conn.cursor()
    try:
        c.execute("INSERT INTO obd2_history (codigo, segmento, montadora, modelo, ano, descricao) VALUES (?, ?, ?, ?, ?, ?)", 
                  (codigo, segmento, montadora, modelo, ano, descricao))
    except sqlite3.OperationalError:
        c.execute("INSERT INTO obd2_history (codigo, montadora, modelo, ano, descricao) VALUES (?, ?, ?, ?, ?)", 
                  (codigo, f"{montadora} ({segmento})", modelo, ano, descricao))
    conn.commit(); conn.close()
    backup_local_para_nuvem()

def carregar_historico_obd2():
    conn = conectar_db()
    try:
        df = pd.read_sql_query("SELECT codigo, segmento, montadora, modelo, ano, data FROM obd2_history ORDER BY id DESC", conn)
    except:
        df = pd.read_sql_query("SELECT codigo, montadora, modelo, ano, data FROM obd2_history ORDER BY id DESC", conn)
    conn.close()
    return df

def diagnostico_avancado_obd2(codigo, segmento="", montadora="", modelo="", ano=""):
    query = f"DTC fault code {codigo}"
    if montadora: query += f" {montadora}"
    if modelo: query += f" {modelo}"
    if segmento and segmento != "Geral / Outros": query += f" {segmento.split('/')[0]}"
    query += " symptoms causes repair troubleshooting"
    
    search_results = ""
    try:
        with DDGS() as ddgs:
            # OTIMIZAÇÃO DE VELOCIDADE: Reduzido para 3 resultados mais relevantes
            resultados = list(ddgs.text(query, max_results=3))
            for r in resultados: search_results += f"- {r['body']}\n"
    except Exception:
        search_results = "(Busca na web indisponível no momento. Use apenas sua base de conhecimento interna.)"
        
    prompt_ia = f"""
    Aja como o 'Chip', um Mecânico Chefe Sênior especialista em múltiplos segmentos (Leves, Pesados, Agrícola, Náutica e Motos).
    O usuário precisa de um laudo TÉCNICO COMPLETO sobre a falha / DTC: **{codigo}**.
    
    Contexto do Veículo/Máquina:
    - Segmento: {segmento or 'Não especificado'}
    - Marca/Montadora: {montadora or 'Qualquer'}
    - Modelo: {modelo or 'Qualquer'}
    - Ano: {ano or 'Qualquer'}
    
    Dados cruzados dos manuais da web agora:
    {search_results}
    
    Gere um relatório técnico de chefe de oficina contendo:
    1. Significado Direto da Falha (DTC).
    2. Descrição Técnica do Manual da Montadora (Descreva com termos técnicos da literatura oficial. Ex: Falha de comunicação, curto ao massa, circuito aberto, tensão acima do limite, etc.).
    3. Aviso Cruzado: Se este código variar dependendo do segmento ou montadora, explique a diferença!
    4. Sintomas percebidos na máquina/veículo.
    5. Causas mais prováveis.
    6. Passos práticos para diagnóstico e solução na oficina.
    Responda em Markdown, com tom profissional, prático e em Português do Brasil.
    """
    
    if HF_TOKEN:
        try:
            client = InferenceClient(token=HF_TOKEN)
            completude = client.chat_completion(
                model="Qwen/Qwen2.5-7B-Instruct",
                messages=[{"role": "user", "content": prompt_ia}],
                max_tokens=850, 
                temperature=0.2 # OTIMIZAÇÃO: Menos "criatividade", mais velocidade e precisão direta
            )
            return completude.choices[0].message.content.strip()
        except Exception as e:
            return f"⚠️ Falha de comunicação no Scanner: {e}\n\n**Dados cruzados:**\n{search_results}"
    else:
        return f"**Dados crus da internet:**\n{search_results}"

# ==========================================
# 5. ROTINAS EEPROM E ARQUIVOS
# ==========================================
def listar_montadoras():
    montadoras = set()
    try:
        conn = conectar_db(); cursor = conn.cursor()
        cursor.execute("SELECT nome FROM montadoras")
        for row in cursor.fetchall(): montadoras.add(higienizar_nome(row[0]))
        conn.close()
    except: pass
    return sorted(list(montadoras))

def listar_modelos(montadora):
    if not montadora: return []
    modelos = set()
    try:
        conn = conectar_db(); cursor = conn.cursor()
        cursor.execute("SELECT modelo FROM veiculos WHERE montadora_nome = ?", (higienizar_nome(montadora),))
        for row in cursor.fetchall(): modelos.add(higienizar_nome(row[0]))
        conn.close()
    except: pass
    return sorted(list(modelos))

def buscar_dados_veiculo_unificado(montadora, modelo):
    try:
        conn = conectar_db(); cursor = conn.cursor()
        cursor.execute("SELECT id, posicao_inicio, intervalo, valores_invertidos, escala, detalhes FROM veiculos WHERE montadora_nome = ? AND modelo = ?", (higienizar_nome(montadora), higienizar_nome(modelo)))
        row = cursor.fetchone()
        if row:
            v_id = row[0]
            cursor.execute("SELECT foto FROM graficos WHERE veiculo_id = ? ORDER BY ordem", (v_id,))
            fotos = [f[0] for f in cursor.fetchall()]
            conn.close()
            return {"id": v_id, "posicao_inicio": row[1], "intervalo": row[2], "valores_invertidos": row[3], "escala": row[4], "detalhes": row[5], "graficos": fotos}
        conn.close()
    except: pass
    return None

def salvar_novo_veiculo_hibrido(montadora, modelo, inicio, intervalo, info_extra, valores_invertidos, escala, imagens_upload=None):
    montadora = higienizar_nome(montadora); modelo = higienizar_nome(modelo)
    conn = conectar_db(); cursor = conn.cursor()
    try:
        cursor.execute("INSERT OR IGNORE INTO montadoras (nome) VALUES (?)", (montadora,))
        cursor.execute("INSERT OR REPLACE INTO veiculos (montadora_nome, modelo, posicao_inicio, intervalo, valores_invertidos, escala, detalhes) VALUES (?, ?, ?, ?, ?, ?, ?)", (montadora, modelo, inicio, intervalo, valores_invertidos, escala, info_extra))
        cursor.execute("SELECT id FROM veiculos WHERE montadora_nome = ? AND modelo = ?", (montadora, modelo))
        veiculo_id = cursor.fetchone()[0]
        pasta_modelo = os.path.join(BASE_DIR, montadora, modelo)
        os.makedirs(pasta_modelo, exist_ok=True)
        with open(os.path.join(pasta_modelo, "dados.json"), "w", encoding="utf-8") as f:
            json.dump({"posicao_inicio": inicio, "intervalo": intervalo, "valores_invertidos": valores_invertidos, "escala": escala, "detalhes": info_extra}, f, indent=4, ensure_ascii=False)
        if imagens_upload:
            cursor.execute("DELETE FROM graficos WHERE veiculo_id = ?", (veiculo_id,))
            for idx, img_file in enumerate(imagens_upload[:6]):
                img_bytes = img_file.read()
                cursor.execute("INSERT INTO graficos (veiculo_id, foto, ordem) VALUES (?, ?, ?)", (veiculo_id, img_bytes, idx+1))
                with open(os.path.join(pasta_modelo, f"grafico_{idx+1}.png"), "wb") as f_img: f_img.write(img_bytes)
        conn.commit(); backup_local_para_nuvem() 
        return True
    except: return False
    finally: conn.close()

def excluir_veiculo_db(montadora, modelo):
    mont = higienizar_nome(montadora); mod = higienizar_nome(modelo)
    conn = conectar_db(); cursor = conn.cursor()
    cursor.execute("DELETE FROM veiculos WHERE montadora_nome = ? AND modelo = ?", (mont, mod))
    conn.commit(); conn.close()
    pasta_modelo = os.path.join(BASE_DIR, mont, mod)
    if os.path.exists(pasta_modelo): shutil.rmtree(pasta_modelo)
    backup_local_para_nuvem() 

def excluir_montadora_db(montadora):
    mont = higienizar_nome(montadora)
    conn = conectar_db(); cursor = conn.cursor()
    cursor.execute("DELETE FROM veiculos WHERE montadora_nome = ?", (mont,))
    cursor.execute("DELETE FROM montadoras WHERE nome = ?", (mont,))
    conn.commit(); conn.close()
    pasta_montadora = os.path.join(BASE_DIR, mont)
    if os.path.exists(pasta_montadora): shutil.rmtree(pasta_montadora)
    backup_local_para_nuvem() 

def obter_resumo_banco_para_ia():
    try:
        conn = conectar_db(); cursor = conn.cursor()
        cursor.execute("SELECT nome FROM montadoras")
        monts = [r[0] for r in cursor.fetchall()]
        cursor.execute("SELECT COUNT(*) FROM obd2_history")
        hist_obd = cursor.fetchone()[0]
        conn.close()
        return f"Montadoras: {', '.join(monts)}\nTotal de Pesquisas OBD-II no Histórico: {hist_obd}"
    except: return "Erro ao carregar dados locais."

# ==========================================
# 6. CORE DE INTELIGÊNCIA ARTIFICIAL (CHIP)
# ==========================================
def processar_linguagem_chip(prompt_cru):
    DADOS_DO_SISTEMA = obter_resumo_banco_para_ia()
    CONTEUDO_DO_SISTEMA = (
        "Você é o Chip, a IA do HyperTork System. Aja como um Mecânico Chefe Sênior de oficina pesada: direto, técnico e com linguajar automotivo real. "
        "Você gerencia mapas de EEPROM e analisa DADOS OBD-II.\n"
        f"Dados locais da nossa oficina: {DADOS_DO_SISTEMA}\n\n"
        "Se o usuário pedir diagnóstico de um código (ex: P0001 na Volvo), escolha a ação DIAGNOSTICAR_FALHA e deixe a resposta vazia.\n"
        "Formato JSON EXATO:\n"
        "{\n"
        '  "acao": "CRIAR_MONTADORA" | "CADASTRAR_VEICULO" | "DIAGNOSTICAR_FALHA" | "ANALISAR_RESPONDER",\n'
        '  "parametros": { "codigo": "P0001", "montadora": "NOME", "modelo": "NOME", "ano": "2014" },\n'
        '  "resposta": "Sua resposta de Mecânico Chefe (para dúvidas ou confirmação)."\n'
        "}\n"
    )
    
    if HF_TOKEN:
        try:
            client = InferenceClient(token=HF_TOKEN)
            completude = client.chat_completion(
                model="Qwen/Qwen2.5-7B-Instruct",
                messages=[{"role": "system", "content": CONTEUDO_DO_SISTEMA}, {"role": "user", "content": prompt_cru}],
                max_tokens=400, temperature=0.2
            )
            resposta_ia = completude.choices[0].message.content.strip()
            resposta_ia = re.sub(r'```json\s*|```', '', resposta_ia)
            dados = json.loads(resposta_ia)
            
            acao = dados.get("acao", "ANALISAR_RESPONDER")
            params = dados.get("parametros", {})
            texto_base = dados.get("resposta", "Entendido, chefe!")
            
            if acao == "DIAGNOSTICAR_FALHA" and params.get("codigo"):
                cod = params.get("codigo")
                mont = params.get("montadora", "")
                mod = params.get("modelo", "")
                ano = params.get("ano", "")
                laudo = diagnostico_avancado_obd2(cod, "", mont, mod, ano)
                salvar_pesquisa_obd2(cod, "IA Chat", mont, mod, ano, laudo)
                return f"🔧 **Laudo de Falha do Chefe:**\n\n{laudo}"
                
            elif acao == "CRIAR_MONTADORA" and params.get("montadora"):
                m = higienizar_nome(params["montadora"])
                conn = conectar_db(); cursor = conn.cursor()
                cursor.execute("INSERT OR IGNORE INTO montadoras (nome) VALUES (?)", (m,))
                conn.commit(); conn.close(); os.makedirs(os.path.join(BASE_DIR, m), exist_ok=True)
                backup_local_para_nuvem()
                return texto_base
                
            return texto_base
        except Exception:
            pass 

    return "🤖 Opa chefe, o servidor da matriz deu uma engasgada. Use os painéis ali do lado enquanto eu reinicio os módulos!"

@st.dialog("🔍 Visualizador de Mapa Ampliado", width="large")
def abrir_modal_zoom(foto_bytes, legenda_titulo):
    st.write(f"#### {legenda_titulo}")
    zoom_dinamico = st.slider("Arraste para ajustar o Zoom do Mapa", 400, 2000, 950, step=50)
    st.image(foto_bytes, width=zoom_dinamico)
    if st.button("❌ Fechar Visualização", use_container_width=True): st.rerun()

# ==========================================
# 7. FUNÇÕES ESPECÍFICAS DA ABA DE GESTÃO DE OS
# ==========================================
def calcular_valor_inicial(linha):
    descricao = str(linha.get("Nome arquivo", "")).upper().strip()
    veiculo = str(linha.get("Fabricante", "")).upper().strip()
    descricao = re.sub(r"\s+", " ", descricao)
    veiculo = re.sub(r"\s+", " ", veiculo)

    termos_stg2 = ["STG2", "STG 2", "STAG2", "STAG 2"]
    termos_mod_off = ["MOD", "OFF"]
    fabricantes_especiais = ["NEW HOLLAND", "VALTRA", "CASE IH", "CASE", "MASSEY FERGUSSON", "MASSEY", "CLAAS", "JHON DEERE", "JOHN DEERE", "DEERE", "FENDT", "JACTO", "DOPPSTADT", "JAN", "VOLVO CONSTRUCTION EQUIPMENT", "VOLVO CONSTRUCTION", "VOLVO CE"]

    eh_especial = any(fab in veiculo for fab in fabricantes_especiais)
    if "VOLVO TRUCK" in veiculo: eh_especial = False

    if any(termo in descricao for termo in termos_stg2): return 1400 if eh_especial else 650
    elif any(termo in descricao for termo in termos_mod_off): return 700 if eh_especial else 350
    return None

def limpar_descricao_os(desc_original):
    desc_upper = str(desc_original).upper().strip()
    if "STAG 2" in desc_upper or "STAG2" in desc_upper: return "STAG 2"
    elif "STG 2" in desc_upper or "STG2" in desc_upper: return "STG 2"
    elif "MOD" in desc_upper: return "MOD"
    elif "OFF" in desc_upper: return "OFF"
    return desc_original

def modificar_modelo_docx(modelo_bytes, flash_point, cliente_nome, cidade, contato, linhas_tabela, total_valor):
    doc = Document(io.BytesIO(modelo_bytes))
    for t in doc.tables:
        for row in t.rows:
            if len(row.cells) >= 2:
                texto_celula_1 = row.cells[0].text.upper().strip()
                if "CLIENTE:" in texto_celula_1:
                    row.cells[1].text = f"{cliente_nome} - {flash_point}"
                    for p in row.cells[1].paragraphs:
                        for run in p.runs: run.font.name = 'Arial'; run.font.size = Pt(11)
                elif "CIDADE:" in texto_celula_1:
                    row.cells[1].text = cidade
                    for p in row.cells[1].paragraphs:
                        for run in p.runs: run.font.name = 'Arial'; run.font.size = Pt(11)
                elif "CONTATO:" in texto_celula_1:
                    row.cells[1].text = contato
                    for p in row.cells[1].paragraphs:
                        for run in p.runs: run.font.name = 'Arial'; run.font.size = Pt(11)

    linhas_validas = [l for l in linhas_tabela if l.get("Valor") is not None and str(l.get("Valor")).strip() != "" and str(l.get("Valor")).lower() != "nan"]
    tabela_servicos = None
    for t in doc.tables:
        if len(t.rows) > 0 and "Nº MAPA" in t.rows[0].cells[0].text.upper():
            tabela_servicos = t
            break
            
    if tabela_servicos:
        for i, linha in enumerate(linhas_validas):
            idx_linha_destino = i + 1  
            if idx_linha_destino >= len(tabela_servicos.rows): row_cells = tabela_servicos.add_row().cells
            else: row_cells = tabela_servicos.rows[idx_linha_destino].cells
                
            dados_linha = [str(linha.get("Nº Mapa", "")), str(linha.get("Data", "")), str(linha.get("Veículo", "")), str(linha.get("Placa", "")), limpar_descricao_os(linha.get("Descrição", "")), f"R$ {linha.get('Valor', '')}"]
            
            for idx_col, valor_celula in enumerate(dados_linha):
                if idx_col < len(row_cells):
                    row_cells[idx_col].text = valor_celula
                    for p_cell in row_cells[idx_col].paragraphs:
                        if idx_col in [0, 1, 3, 5]: p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p_cell.runs: r.font.name = 'Arial'; r.font.size = Pt(10)
                            
        linha_inicio_remocao = len(linhas_validas) + 1
        while len(tabela_servicos.rows) > linha_inicio_remocao:
            linha_para_apagar = tabela_servicos.rows[linha_inicio_remocao]
            tabela_servicos._tbl.remove(linha_para_apagar._tr)

    if pd.isna(total_valor) or str(total_valor).lower() == "nan": total_valor = 0.0
    valor_formatado_texto = f"{float(total_valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    
    for t in doc.tables:
        for row in t.rows:
            contem_total = any("TOTAL" in cell.text.upper() for cell in row.cells)
            if contem_total:
                for cell in row.cells:
                    if "R$" in cell.text or "NAN" in cell.text.upper() or cell.text.strip() == "":
                        cell.text = f"R$ {valor_formatado_texto}"
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in p.runs:
                                run.bold = True
                                run.font.name = 'Arial'
                                run.font.size = Pt(12)
                                run.font.color.rgb = RGBColor(234, 88, 12)

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

# ==========================================
# 8. ESTILIZAÇÃO CSS E RESPONSIVIDADE
# ==========================================
st.markdown("""
    <style>
    /* Estabilidade de Tela e Correções Gerais */
    html, body, [class*="css"]  { overflow-x: hidden; }
    .block-container { padding-top: 2rem; max-width: 1200px; }
    
    /* -----------------------------------------------------------
       ESTILIZAÇÃO DOS BOTÕES GIGANTES (HTML PURO)
    ----------------------------------------------------------- */
    .big-hub-btn-link {
        text-decoration: none !important;
        display: block !important;
        width: 100% !important;
    }
    
    .big-hub-btn {
        display: flex !important;
        flex-direction: column !important;
        align-items: center !important;
        justify-content: center !important;
        height: 150px !important; /* ALTURA TRAVADA E FIXA */
        min-height: 150px !important;
        max-height: 150px !important;
        padding: 15px !important; 
        box-sizing: border-box !important; 
        color: white !important;
        border-radius: 15px !important;
        box-shadow: 0 4px 10px rgba(0,0,0,0.3) !important;
        cursor: pointer !important;
        transition: transform 0.2s, box-shadow 0.2s !important;
        text-align: center !important;
        margin-bottom: 20px !important;
        overflow: hidden !important; 
    }
    
    .btn-blue { background: linear-gradient(145deg, #1E88E5, #1565C0); }
    .btn-red { background: linear-gradient(145deg, #E53935, #C62828); }
    .btn-green { background: linear-gradient(145deg, #43A047, #2E7D32); }
    
    .big-hub-btn:hover {
        transform: translateY(-5px);
        box-shadow: 0 12px 24px rgba(0,0,0,0.4) !important;
    }
    
    .big-hub-btn h2 { color: white !important; margin: 8px 0 0 0 !important; font-weight: bold; font-size: 1.2rem !important; line-height: 1.2 !important; word-wrap: break-word; text-transform: uppercase;}
    .big-hub-btn span { font-size: 2.5rem !important; line-height: 1 !important;}

    /* -----------------------------------------------------------
       ESTILIZAÇÃO DO CHIP (BOTÃO LARANJA FLUTUANTE)
    ----------------------------------------------------------- */
    div[data-testid="stPopover"] {
        position: fixed !important;
        bottom: 20px !important;
        right: 20px !important;
        z-index: 999999 !important;
        width: 70px !important;
        height: 70px !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }
    
    div[data-testid="stPopover"] > button, 
    div[data-testid="stPopover"] > div > button {
        background-color: #FF8C00 !important;
        border: 2px solid #E67E22 !important;
        border-radius: 50% !important;
        width: 65px !important;
        height: 65px !important;
        min-width: 65px !important;
        min-height: 65px !important;
        box-shadow: 0 8px 16px rgba(255, 140, 0, 0.4) !important;
        transition: transform 0.2s !important;
        padding: 0 !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
    }
    
    div[data-testid="stPopover"] > button:hover,
    div[data-testid="stPopover"] > div > button:hover {
        transform: scale(1.1) !important;
    }
    
    div[data-testid="stPopover"] > button p, 
    div[data-testid="stPopover"] > div > button p {
        font-size: 32px !important;
        line-height: 1 !important;
        margin: 0 !important;
        display: block !important;
    }
    
    div[data-testid="stPopoverBody"] button {
        background-color: transparent !important;
        border: 1px solid #ddd !important;
        border-radius: 8px !important;
        width: auto !important;
        height: auto !important;
        min-width: 0 !important;
        min-height: 0 !important;
        box-shadow: none !important;
    }
    div[data-testid="stPopoverBody"] div[data-testid="stForm"] button {
        background-color: #1E88E5 !important;
        color: white !important;
        border: none !important;
    }

    div[data-testid="stPopoverBody"] {
        width: 90vw !important;
        max-width: 380px !important;
        border-radius: 15px !important;
        box-shadow: 0 10px 30px rgba(0,0,0,0.3) !important;
        padding: 1rem !important;
        position: fixed !important;
        right: 20px !important;
        bottom: 100px !important;
        top: auto !important;
        left: auto !important;
        transform: none !important;
    }
    </style>
""", unsafe_allow_html=True)

# Lógica de Roteamento baseada em Parâmetros de URL
params = st.query_params
if "page" in params:
    st.session_state.app_mode = params["page"]
elif 'app_mode' not in st.session_state:
    st.session_state.app_mode = "HOME"

if 'montadora_selecionada' not in st.session_state:
    st.session_state.montadora_selecionada = ""
if 'chat_historico' not in st.session_state:
    st.session_state.chat_historico = [{"role": "assistant", "content": "Oi! Eu sou o Chip, como posso ajudar?"}]
if "df_filtrado" not in st.session_state:
    st.session_state.df_filtrado = None

# ==========================================
# 9. BARRA LATERAL E NAVEGAÇÃO
# ==========================================
st.sidebar.title("🛡️ HyperTork Hub")
if st.session_state.app_mode != "HOME":
    if st.sidebar.button("🎮 Voltar ao Menu Principal", use_container_width=True, type="primary"):
        st.query_params.clear()
        st.session_state.app_mode = "HOME"
        st.session_state.montadora_selecionada = ""
        st.rerun()
else:
    st.sidebar.info("📌 Escolha uma das ferramentas ao lado.")

montadoras_existentes = listar_montadoras()

# ==========================================
# 10. RENDERIZAÇÃO DAS TELAS
# ==========================================
if st.session_state.app_mode == "HOME":
    
    caminho_logo = os.path.join(LOGOS_DIR, "logo.png")
    if os.path.exists(caminho_logo):
        col_logo1, col_logo2, col_logo3 = st.columns([1, 2, 1])
        with col_logo2:
            try:
                st.image(caminho_logo, use_container_width=True)
            except Exception:
                pass
        st.markdown("---")

    st.markdown("<h1 style='text-align: center; margin-bottom: 50px;'>HyperTork System Hub</h1>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown(f"""
            <a href="?page=EEPROM" class="big-hub-btn-link" target="_self">
                <div class="big-hub-btn btn-blue">
                    <span>⚙️</span>
                    <h2>Gráficos EEPROM</h2>
                </div>
            </a>
        """, unsafe_allow_html=True)
            
    with col2:
        st.markdown(f"""
            <a href="?page=OBD2" class="big-hub-btn-link" target="_self">
                <div class="big-hub-btn btn-red">
                    <span>🚗</span>
                    <h2>Códigos de Falha</h2>
                </div>
            </a>
        """, unsafe_allow_html=True)
        
    with col3:
        st.markdown(f"""
            <a href="?page=GESTAO_OS" class="big-hub-btn-link" target="_self">
                <div class="big-hub-btn btn-green">
                    <span>📊</span>
                    <h2>Gestão & OS</h2>
                </div>
            </a>
        """, unsafe_allow_html=True)

# ------------------------------------------
# TELA 1: OBD-II / DTC SCANNERS (BUSCA UNIVERSAL)
# ------------------------------------------
elif st.session_state.app_mode == "OBD2":
    st.title("🛠️ Diagnóstico Universal de Falhas (DTC)")
    st.markdown("Busca independente avançada para qualquer veículo ou máquina. **Não necessita de cadastro na aba EEPROM.**")
    
    with st.container(border=True):
        col_cod, col_seg, col_mont, col_mod, col_ano = st.columns([2, 2, 2, 2, 1])
        
        with col_cod:
            codigo_input = st.text_input("Código de Falha", placeholder="Ex: P0001, SPN 3216").strip().upper()
            
        with col_seg:
            opcoes_segmento = [
                "Geral / Outros", 
                "Pesado / Caminhão / Ônibus", 
                "Leve / Carro de Passeio", 
                "Agrícola / Trator", 
                "Construção / Linha Amarela", 
                "Náutica / Marítimo", 
                "Motos / Quadriciclos"
            ]
            segmento_input = st.selectbox("Segmento", opcoes_segmento)
            
        with col_mont:
            mont_input = st.text_input("Marca / Montadora", placeholder="Ex: Volvo, John Deere...").strip()
            
        with col_mod:
            mod_input = st.text_input("Modelo", placeholder="Ex: FH 460, 8R 340...").strip()
            
        with col_ano:
            ano_input = st.text_input("Ano", placeholder="Ex: 2020").strip()
            
        btn_buscar = st.button("🔍 Iniciar Diagnóstico Universal", use_container_width=True, type="primary")
            
    if btn_buscar:
        if codigo_input:
            with st.spinner(f"Cruzando bancos de dados e manuais de {segmento_input} para {codigo_input}..."):
                descricao_encontrada = diagnostico_avancado_obd2(codigo_input, segmento_input, mont_input, mod_input, ano_input)
                st.subheader(f"Laudo Técnico: {codigo_input} | {mont_input or 'Marca Não Informada'}")
                st.info(descricao_encontrada)
                salvar_pesquisa_obd2(codigo_input, segmento_input, mont_input, mod_input, ano_input, descricao_encontrada)
                st.success("✅ Diagnóstico salvo no histórico da nuvem!")
        else:
            st.warning("O código da falha é obrigatório para iniciar o diagnóstico.")
            
    st.divider()
    st.subheader("📚 Histórico de Pesquisas")
    df_historico = carregar_historico_obd2()
    if not df_historico.empty:
        st.dataframe(df_historico, use_container_width=True, hide_index=True)
    else:
        st.write("Nenhuma falha foi pesquisada ainda.")

# ------------------------------------------
# TELA 2: GESTÃO EEPROM
# ------------------------------------------
elif st.session_state.app_mode == "EEPROM":
    if st.session_state.montadora_selecionada == "":
        st.title("🚜 Painel de Controle - Baias EEPROM")
        st.markdown("### Escolha a Montadora desejada para abrir os modelos")
        st.write("")
        if not montadoras_existentes:
            st.info("Nenhuma montadora cadastrada. Fale com o Chip ali no canto inferior direito para criar uma!")
        else:
            cols = st.columns(4)
            for i, m in enumerate(montadoras_existentes):
                with cols[i % 4]:
                    with st.container(border=True):
                        caminho_logo = buscar_logo_montadora_automatica(m)
                        if caminho_logo:
                            logo_html_src = obter_image_base64_html(caminho_logo)
                            if logo_html_src:
                                st.markdown(f"""
                                    <div style="display: flex; justify-content: center; align-items: center; 
                                                background-color: #262730; padding: 10px; border-radius: 8px; 
                                                height: 110px; width: 100%; box-sizing: border-box; margin-bottom: 6px;
                                                box-shadow: inset 0 0 5px rgba(0,0,0,0.5);">
                                        <img src="{logo_html_src}" style="max-height: 90px; max-width: 100%; object-fit: contain; filter: drop-shadow(2px 2px 4px rgba(0,0,0,0.6));">
                                    </div>
                                """, unsafe_allow_html=True)
                        else:
                            st.markdown(f"""
                                <div style="display: flex; justify-content: center; align-items: center; height: 110px; width: 100%; margin-bottom: 6px; background-color: #262730; border-radius: 8px; box-shadow: inset 0 0 5px rgba(0,0,0,0.5);">
                                    <p style='text-align:center; font-weight:bold; color:#1E88E5; margin:0;'>🏭 {m}</p>
                                </div>
                            """, unsafe_allow_html=True)
                        if st.button(f"Abrir {m}", key=f"home_{m}", use_container_width=True):
                            st.session_state.montadora_selecionada = m
                            st.rerun()
    else:
        col_logo, col_nome = st.columns([1, 8])
        caminho_da_logo = buscar_logo_montadora_automatica(st.session_state.montadora_selecionada)
        with col_logo:
            if caminho_da_logo:
                logo_html_src = obter_image_base64_html(caminho_da_logo)
                if logo_html_src:
                    st.markdown(f"""
                        <div style="display: flex; justify-content: center; align-items: center; 
                                    background-color: #262730; padding: 6px; border-radius: 8px; height: 75px; width: 75px; box-sizing: border-box;
                                    box-shadow: inset 0 0 5px rgba(0,0,0,0.5);">
                            <img src="{logo_html_src}" style="max-height: 60px; max-width: 100%; object-fit: contain; filter: drop-shadow(1px 1px 2px rgba(0,0,0,0.5));">
                        </div>
                    """, unsafe_allow_html=True)
                else: st.subheader("🏭")
            else: st.subheader("🏭")
                
        with col_nome:
            st.markdown(f"<h1 style='margin-top: 5px; color: #1E88E5;'>{st.session_state.montadora_selecionada}</h1>", unsafe_allow_html=True)
        st.markdown("---")

        modelos_existentes = listar_modelos(st.session_state.montadora_selecionada)
        if not modelos_existentes:
            st.warning(f"Nenhum veículo cadastrado para a montadora {st.session_state.montadora_selecionada}.")
        else:
            escolha_modelo = st.selectbox("📂 Selecione o Veículo para carregar os gráficos:", [""] + modelos_existentes)
            st.write("")
            if escolha_modelo:
                st.markdown(f"#### 📍 Mapa: {st.session_state.montadora_selecionada} {escolha_modelo}")
                dados_mapa = buscar_dados_veiculo_unificado(st.session_state.montadora_selecionada, escolha_modelo)

                col_info, col_img = st.columns([1, 2])
                
                with col_info:
                    with st.container(border=True):
                        st.subheader("📋 Informações Gerais")
                        if dados_mapa:
                            st.write("**Início do Gráfico:**")
                            st.code(dados_mapa["posicao_inicio"], language="text")
                            st.write("**Intervalo de Endereços:**")
                            st.code(dados_mapa["intervalo"], language="text")
                            st.write("**Detalhes do Veículo:**")
                            st.info(dados_mapa["detalhes"])
                    
                    with st.container(border=True):
                        st.markdown("⚙️ **Configuração de Mapa**")
                        if dados_mapa:
                            st.write(f"**Valores invertidos:** {dados_mapa.get('valores_invertidos', 'Desativado')}")
                            st.write(f"**Escala:** {dados_mapa.get('escala', '8 bits')}")

                with col_img:
                    if not dados_mapa or not dados_mapa["graficos"]:
                        st.error("⚠️ Nenhuma imagem de mapa encontrada para este veículo.")
                    else:
                        st.caption("💡 *Dica:* Clique nos botões de Expandir para ver os mapas detalhados em Tela Cheia!")
                        lista_fotos = dados_mapa["graficos"]
                        for idx in range(0, len(lista_fotos), 2):
                            sub_cols = st.columns(2)
                            with sub_cols[0]:
                                if idx < len(lista_fotos):
                                    label_cap = f"Gráfico Principal ({idx+1})"
                                    st.image(lista_fotos[idx], use_container_width=True, caption=label_cap)
                                    if st.button(f"🔍 Expandir ({idx+1})", key=f"btn_zoom_{idx}", use_container_width=True):
                                        abrir_modal_zoom(lista_fotos[idx], label_cap)
                            with sub_cols[1]:
                                if idx + 1 < len(lista_fotos):
                                    label_cap_2 = f"Gráfico Complementar ({idx+2})"
                                    st.image(lista_fotos[idx+1], use_container_width=True, caption=label_cap_2)
                                    if st.button(f"🔍 Expandir ({idx+2})", key=f"btn_zoom_{idx+1}", use_container_width=True):
                                        abrir_modal_zoom(lista_fotos[idx+1], label_cap_2)
                    
    st.markdown("<br><br>", unsafe_allow_html=True)

    with st.expander("➕ CADASTRAR: Adicionar Estruturas Independentes"):
        cad_tab1, cad_tab2 = st.tabs(["🏭 Cadastrar Montadora", "🚗 Cadastrar Veículo"])
        with cad_tab1:
            st.subheader("Nova Montadora")
            nova_m = st.text_input("Digite o Nome da Montadora", key="input_nova_m").strip()
            if st.button("Efetivar Montadora", type="primary"):
                if not nova_m: st.error("❌ O campo não pode ficar em branco!")
                else:
                    m_hig = higienizar_nome(nova_m)
                    if m_hig in montadoras_existentes: st.error(f"❌ Montadora '{m_hig}' já cadastrada!")
                    else:
                        conn = conectar_db(); cursor = conn.cursor()
                        cursor.execute("INSERT INTO montadoras (nome) VALUES (?)", (m_hig,))
                        conn.commit(); conn.close()
                        os.makedirs(os.path.join(BASE_DIR, m_hig), exist_ok=True)
                        backup_local_para_nuvem()
                        st.success(f"✅ Montadora '{m_hig}' salva no cofre!")
                        st.rerun()
                        
        with cad_tab2:
            st.subheader("Novo Veículo")
            if not montadoras_existentes: st.info("Cadastre ao menos uma montadora para liberar.")
            else:
                m_form = st.selectbox("Selecione a Montadora Alvo", montadoras_existentes, key="sb_m_form")
                v_form = st.text_input("Nome do Modelo / Veículo", key="input_v_form").strip()
                vc1, vc2 = st.columns(2)
                v_ini = vc1.text_input("Endereço Inicial", key="ini_v_form")
                v_int = vc2.text_input("Intervalo de Endereço", key="int_v_form")
                v_inv = st.selectbox("Valores Invertidos?", ["Desativado", "Ativado"], key="inv_v_form")
                v_esc = st.selectbox("Escala do Mapa", ["8 bits", "16 bits", "32 bits"], key="esc_v_form")
                v_det = st.text_area("Detalhes Adicionais", key="det_v_form")
                v_files = st.file_uploader("Fotos dos Gráficos (Máx 6)", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="files_v_form")
                if st.button("Efetivar Veículo", type="primary"):
                    v_hig = higienizar_nome(v_form)
                    if not v_form: st.error("❌ O nome do veículo é obrigatório!")
                    else:
                        modelos_da_marca = listar_modelos(m_form)
                        if v_hig in modelos_da_marca: st.error(f"❌ Veículo '{v_hig}' já existe na {m_form}!")
                        else:
                            status_save = salvar_novo_veiculo_hibrido(m_form, v_form, v_ini, v_int, v_det, v_inv, v_esc, v_files)
                            if status_save:
                                st.success(f"✅ Ficha técnica '{v_hig}' sincronizada!")
                                st.rerun()

    with st.expander("⚙️ GERENCIAR: Painel de Edição e Exclusão Total"):
        ger_tab1, ger_tab2 = st.tabs(["🏭 Gerenciar Montadoras", "🚗 Gerenciar Veículos"])
        with ger_tab1:
            st.subheader("Modificação de Marcas")
            if not montadoras_existentes: st.warning("Nenhuma montadora localizada.")
            else:
                m_select_edit = st.selectbox("Escolha a Montadora para Alterar", montadoras_existentes, key="sb_m_edit_pane")
                novo_nome_m = st.text_input("Alterar Nome da Montadora para:", value=m_select_edit, key="txt_rename_m").strip()
                m_ed_col1, m_ed_col2 = st.columns(2)
                if m_ed_col1.button("💾 Salvar Novo Nome", key="btn_rename_m"):
                    n_m_hig = higienizar_nome(novo_nome_m)
                    if not n_m_hig: st.error("❌ Nome inválido.")
                    elif n_m_hig in montadoras_existentes and n_m_hig != m_select_edit: st.error("❌ Já existe!")
                    else:
                        conn = conectar_db(); cursor = conn.cursor()
                        cursor.execute("UPDATE montadoras SET nome = ? WHERE nome = ?", (n_m_hig, m_select_edit))
                        cursor.execute("UPDATE veiculos SET montadora_nome = ? WHERE montadora_nome = ?", (n_m_hig, m_select_edit))
                        conn.commit(); conn.close()
                        old_path = os.path.join(BASE_DIR, m_select_edit)
                        new_path = os.path.join(BASE_DIR, n_m_hig)
                        if os.path.exists(old_path): os.rename(old_path, new_path)
                        backup_local_para_nuvem()
                        st.success("✅ Nome atualizado globalmente!")
                        st.session_state.montadora_selecionada = ""
                        st.rerun()
                if m_ed_col2.button("🗑️ Excluir Montadora", key="btn_del_m_pane"):
                    excluir_montadora_db(m_select_edit)
                    st.success(f"✅ Limpeza concluída e Nuvem sincronizada.")
                    st.session_state.montadora_selecionada = ""
                    st.rerun()

        with ger_tab2:
            st.subheader("Edição Completa de Veículos")
            if not montadoras_existentes: st.warning("Sem marcas salvas.")
            else:
                m_sel_v = st.selectbox("Filtrar por Montadora", montadoras_existentes, key="sb_m_sel_v")
                v_existentes = listar_modelos(m_sel_v)
                if not v_existentes: st.info("Nenhum veículo localizado nesta marca.")
                else:
                    v_sel_edit = st.selectbox("Selecione o Veículo", v_existentes, key="sb_v_sel_edit")
                    dados_v = buscar_dados_veiculo_unificado(m_sel_v, v_sel_edit)
                    if dados_v:
                        st.markdown("---")
                        v_novo_nome = st.text_input("Alterar Nome do Veículo", value=v_sel_edit, key="txt_v_name_edit").strip()
                        ve_c1, ve_c2 = st.columns(2)
                        v_novo_ini = ve_c1.text_input("Alterar Endereço Inicial", value=dados_v["posicao_inicio"], key="txt_v_ini_edit")
                        v_novo_int = ve_c2.text_input("Alterar Intervalo", value=dados_v["intervalo"], key="txt_v_int_edit")
                        ve_c3, ve_c4 = st.columns(2)
                        inv_idx = 0 if dados_v["valores_invertidos"] == "Desativado" else 1
                        v_novo_inv = ve_c3.selectbox("Valores Invertidos?", ["Desativado", "Ativado"], index=inv_idx, key="sb_v_inv_edit")
                        esc_opcoes = ["8 bits", "16 bits", "32 bits"]
                        esc_idx = esc_opcoes.index(dados_v["escala"]) if dados_v["escala"] in esc_opcoes else 0
                        v_novo_esc = ve_c4.selectbox("Escala do Mapa", esc_opcoes, index=esc_idx, key="sb_v_esc_edit")
                        v_novo_det = st.text_area("Alterar Detalhes do Veículo", value=dados_v["detalhes"], key="txt_v_det_edit")
                        st.warning("⚠️ Novas fotos substituirão todas as imagens antigas!")
                        v_novas_fotos = st.file_uploader("Substituir Imagens (Máx 6)", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="files_v_edit")
                        v_manage_col1, v_manage_col2 = st.columns(2)
                        if v_manage_col1.button("💾 Salvar Alterações", type="primary", key="btn_save_v_edit"):
                            n_v_hig = higienizar_nome(v_novo_nome)
                            if not n_v_hig: st.error("❌ O nome não pode ser vazio.")
                            else:
                                if n_v_hig != v_sel_edit and n_v_hig in v_existentes: st.error("❌ Já existe outro modelo com esse nome!")
                                else:
                                    if n_v_hig != v_sel_edit:
                                        velha_pasta = os.path.join(BASE_DIR, m_sel_v, v_sel_edit)
                                        if os.path.exists(velha_pasta): shutil.rmtree(velha_pasta)
                                        conn = conectar_db(); cursor = conn.cursor()
                                        cursor.execute("UPDATE veiculos SET modelo = ? WHERE id = ?", (n_v_hig, dados_v["id"]))
                                        conn.commit(); conn.close()
                                    if n_v_hig != v_sel_edit and not v_novas_fotos:
                                        salvar_novo_veiculo_hibrido(m_sel_v, n_v_hig, v_novo_ini, v_novo_int, v_novo_det, v_novo_inv, v_novo_esc, None)
                                        sincronizar_banco_com_pastas()
                                    else:
                                        salvar_novo_veiculo_hibrido(m_sel_v, n_v_hig, v_novo_ini, v_novo_int, v_novo_det, v_novo_inv, v_novo_esc, v_novas_fotos)
                                    st.success(f"✅ Sucesso: Sincronizado com estabilidade!")
                                    st.rerun()
                        if v_manage_col2.button("🗑️ Excluir Veículo", key="btn_del_v_edit"):
                            excluir_veiculo_db(m_sel_v, v_sel_edit)
                            st.success(f"✅ Remoção Concluída!")
                            st.rerun()

# ------------------------------------------
# TELA 3: GESTÃO DE SERVIÇOS E OS
# ------------------------------------------
elif st.session_state.app_mode == "GESTAO_OS":
    st.title("📊 Gestão de Serviços & Emissão de OS")
    st.write("Filtragem, cálculo de valores, remoção de duplicadas por Matrícula e preenchimento automático do modelo Word da Hyper Tork.")

    aba1, aba2 = st.tabs(["📋 Processamento da Planilha", "📄 Gerar Ordem de Serviço"])

    with aba1:
        arquivo_carregado = st.file_uploader(
            "Arraste ou selecione a planilha FPF_List para iniciar:",
            type=["xlsx", "xls", "csv"],
            key="uploader_planilha"
        )

        if arquivo_carregado is not None:
            try:
                conteudo = arquivo_carregado.read()
                try:
                    excel_file = pd.ExcelFile(io.BytesIO(conteudo))
                    abas = excel_file.sheet_names
                    if len(abas) > 1:
                        aba_selecionada = st.selectbox("Selecione a aba com os dados:", abas, key="selecao_abas_app")
                    else:
                        aba_selecionada = abas[0]
                    df = pd.read_excel(io.BytesIO(conteudo), sheet_name=aba_selecionada)
                except Exception:
                    try:
                        df = pd.read_csv(io.BytesIO(conteudo), sep=";", encoding="utf-8")
                        if df.shape[1] <= 1:
                            df = pd.read_csv(io.BytesIO(conteudo), sep=",", encoding="utf-8")
                    except Exception:
                        df = pd.read_csv(io.BytesIO(conteudo), sep=";", encoding="iso-8859-1")

                if df is None or df.empty or len(df.columns) == 0:
                    st.error("Erro: Não foi possível processar a estrutura de dados deste arquivo.")
                else:
                    df.columns = df.columns.str.strip()

                    if "T" in df.columns:
                        df["T"] = df["T"].astype(str).str.strip()
                        df_filtrado = df[df["T"] == "MOD"].copy()
                    else:
                        st.warning("Aviso: A coluna 'T' não foi encontrada.")
                        df_filtrado = df.copy()

                    colunas_originais = ["Arquivo ID", "Fabricante", "Matrícula", "FlashPoint", "Cliente", "Nome arquivo", "Dada"]
                    colunas_existentes = [col for col in colunas_originais if col in df_filtrado.columns]
                    df_filtrado = df_filtrado[colunas_existentes].copy()

                    df_filtrado["Valor"] = df_filtrado.apply(calcular_valor_inicial, axis=1)

                    dicionario_renomear = {
                        "Arquivo ID": "Nº Mapa",
                        "Fabricante": "Veículo",
                        "Matrícula": "Placa",
                        "Nome arquivo": "Descrição",
                        "Dada": "Data",
                        "FlashPoint": "Flash Point",
                    }
                    df_filtrado = df_filtrado.rename(columns=dicionario_renomear)

                    ordem_solicitada = ["Nº Mapa", "Data", "Veículo", "Placa", "Flash Point", "Cliente", "Descrição", "Valor"]
                    colunas_finais = [col for col in ordem_solicitada if col in df_filtrado.columns]
                    df_filtrado = df_filtrado[colunas_finais].copy()

                    if "Data" in df_filtrado.columns:
                        df_filtrado["Data"] = pd.to_datetime(df_filtrado["Data"], errors='coerce')
                        df_filtrado["Data"] = df_filtrado["Data"].dt.strftime('%d/%m/%Y').fillna("")

                    if "Flash Point" in df_filtrado.columns:
                        df_filtrado["Flash Point"] = df_filtrado["Flash Point"].astype(str).str.strip()
                        df_filtrado = df_filtrado.sort_values(by=["Flash Point", "Nº Mapa"] if "Nº Mapa" in df_filtrado.columns else ["Flash Point"], ascending=True)

                    st.session_state.df_filtrado = df_filtrado.copy()

                    if not df_filtrado.empty and "Flash Point" in df_filtrado.columns:
                        lista_linhas = []
                        linhas_amarelas = []
                        linhas_laranjas = []
                        contador_linha_excel = 2

                        for fp, bloco in df_filtrado.groupby("Flash Point", sort=False):
                            placas_vistas = set()

                            for idx, linha in bloco.iterrows():
                                linha_dict = linha.to_dict()
                                placa_atual = str(linha.get("Placa", "")).strip()

                                if placa_atual in placas_vistas and placa_atual != "":
                                    linha_dict["Valor"] = None
                                    linhas_amarelas.append(contador_linha_excel)
                                else:
                                    if placa_atual != "":
                                        placas_vistas.add(placa_atual)

                                lista_linhas.append(linha_dict)
                                contador_linha_excel += 1

                            df_bloco_temp = pd.DataFrame(lista_linhas[-len(bloco) :])
                            soma_bloco = pd.to_numeric(df_bloco_temp["Valor"], errors="coerce").sum()

                            linha_total = {col: "" for col in colunas_finais}
                            linha_total["Flash Point"] = fp
                            if "Cliente" in linha_total:
                                linha_total["Cliente"] = str(bloco.iloc[0].get("Cliente", ""))
                            linha_total["Descrição"] = "VALOR TOTAL:"
                            linha_total["Valor"] = float(soma_bloco) if soma_bloco > 0 else ""

                            lista_linhas.append(linha_total)
                            linhas_laranjas.append(contador_linha_excel)
                            contador_linha_excel += 1

                            linha_espacamento = {col: "" for col in colunas_finais}
                            lista_linhas.append(linha_espacamento)
                            contador_linha_excel += 1

                        if lista_linhas:
                            lista_linhas.pop()

                        df_excel_final = pd.DataFrame(lista_linhas, columns=colunas_finais)

                        st.subheader("📋 Visualização Prévia dos Dados Processados")
                        st.dataframe(df_filtrado)

                        buffer = io.BytesIO()
                        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
                            df_excel_final.to_excel(writer, index=False, sheet_name="FPF Realizados")

                            workbook = writer.book
                            worksheet = writer.sheets["FPF Realizados"]

                            amarelo_claro = PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid")
                            laranja_claro = PatternFill(start_color="FFE6CC", end_color="FFE6CC", fill_type="solid")

                            for num_linha in linhas_amarelas:
                                for col_idx in range(1, len(colunas_finais) + 1):
                                    worksheet.cell(row=num_linha, column=col_idx).fill = amarelo_claro

                            for num_linha in linhas_laranjas:
                                for col_idx in range(1, len(colunas_finais) + 1):
                                    worksheet.cell(row=num_linha, column=col_idx).fill = laranja_claro

                        st.success("Planilha processada com sucesso na memória! Vá para a aba ao lado para gerar Ordens de Serviço.")
                        st.download_button(
                            label="📥 Baixar Planilha Processada (Excel)",
                            data=buffer.getvalue(),
                            file_name="FPF_Relatorio_Final.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        )
            except Exception as e:
                st.error(f"Erro crítico no processamento: {e}")

    with aba2:
        st.subheader("📄 Emissor de Ordem de Serviço com Base no Modelo Original")
        
        modelo_word_carregado = st.file_uploader(
            "Selecione o seu arquivo original 'MODELO - HYPER TORK PERFORMANCE.docx':",
            type=["docx"],
            key="uploader_modelo_word"
        )
        
        if st.session_state.df_filtrado is None or st.session_state.df_filtrado.empty:
            st.info("Aguardando o upload e processamento da planilha FPF_List na primeira aba para liberar o emissor.")
        elif modelo_word_carregado is None:
            st.warning("Por favor, anexe o arquivo original do seu modelo acima para habilitar o preenchimento automático.")
        else:
            df_base_os = st.session_state.df_filtrado
            bytes_modelo = modelo_word_carregado.read()
            
            lista_fp_unicos = sorted(list(set(str(val).strip() for val in df_base_os["Flash Point"].unique() if pd.notna(val))))
            
            fp_selecionado = st.selectbox("Selecione o Flash Point para gerar a OS correspondente:", lista_fp_unicos)
            
            dados_bloco = df_base_os[df_base_os["Flash Point"] == fp_selecionado]
            cliente_sugerido = str(dados_bloco.iloc[0].get("Cliente", "Cliente Não Identificado"))
            
            col1, col2 = st.columns(2)
            with col1:
                nome_cliente_input = st.text_input("Cliente (Preenchido Automaticamente):", value=cliente_sugerido)
                cidade_input = st.text_input("Cidade (Adicionar a critério do usuário):", placeholder="Ex: Cascavel - PR")
            with col2:
                flash_point_confirmacao = st.text_input("Flash Point Relacionado:", value=fp_selecionado, disabled=True)
                contato_input = st.text_input("Contato (Adicionar a critério do usuário):", placeholder="Ex: (45) 99999-9999")
                
            st.write("### Serviços com Valores Definidos que farão parte desta OS (Linhas vazias serão excluídas do Word):")
            
            linhas_os_finais = []
            placas_vistas_os = set()
            soma_total_os = 0
            
            for idx, row in dados_bloco.iterrows():
                row_dict = row.to_dict()
                placa = str(row_dict.get("Placa", "")).strip()
                
                if placa in placas_vistas_os and placa != "":
                    row_dict["Valor"] = None
                else:
                    if placa != "":
                        placas_vistas_os.add(placa)
                    if row_dict["Valor"] is not None and str(row_dict["Valor"]).lower() != "nan" and not pd.isna(row_dict["Valor"]):
                        soma_total_os += float(row_dict["Valor"])
                
                linhas_os_finais.append(row_dict)
                
            df_preview_os = pd.DataFrame(linhas_os_finais)
            df_preview_os = df_preview_os[df_preview_os["Valor"].notna()].copy()
            df_preview_os["Descrição"] = df_preview_os["Descrição"].apply(limpar_descricao_os)
            
            colunas_preview_os = ["Nº Mapa", "Data", "Veículo", "Placa", "Descrição", "Valor"]
            colunas_preview_existentes = [c for c in colunas_preview_os if c in df_preview_os.columns]
            st.dataframe(df_preview_os[colunas_preview_existentes])
            
            st.metric(label="Valor Total Consolidado da OS (Apenas linhas válidas)", value=f"R$ {soma_total_os:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."))
            
            if st.button("🚀 Preencher e Gerar Ordem de Serviço"):
                arquivo_word_final = modificar_modelo_docx(
                    modelo_bytes=bytes_modelo,
                    flash_point=fp_selecionado,
                    cliente_nome=nome_cliente_input,
                    cidade=cidade_input,
                    contato=contato_input,
                    linhas_tabela=linhas_os_finais,
                    total_valor=soma_total_os
                )
                
                st.success(f"Ordem de Serviço para o Flash Point {fp_selecionado} gerada com sucesso!")
                
                st.download_button(
                    label="📥 Baixar Ordem de Serviço Pronta (.docx)",
                    data=arquivo_word_final.getvalue(),
                    file_name=f"OS_Hyper_Tork_{fp_selecionado}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# ==========================================
# 11. CHIP POP-UP FLUTUANTE (BOTÃO LARANJA)
# ==========================================
with st.popover("🤖"):
    st.markdown("#### 💬 Chip - IA")
    
    chat_container = st.container(height=350)
    with chat_container:
        for msg in st.session_state.chat_historico:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                
    with st.form("chip_chat_form", clear_on_submit=True):
        cols = st.columns([7, 3])
        with cols[0]:
            prompt = st.text_input("Mensagem", label_visibility="collapsed", placeholder="Digite algo...")
        with cols[1]:
            submit = st.form_submit_button("Enviar", use_container_width=True)
        
        if submit and prompt.strip():
            if prompt.strip().lower() in ["/limpar", "limpar chat"]:
                st.session_state.chat_historico = [{"role": "assistant", "content": "Oi! Eu sou o Chip, como posso ajudar?"}]
                st.rerun()
            else:
                st.session_state.chat_historico.append({"role": "user", "content": prompt})
                resposta = processar_linguagem_chip(prompt)
                st.session_state.chat_historico.append({"role": "assistant", "content": resposta})
                st.rerun()